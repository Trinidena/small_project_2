import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.cluster import KMeans, DBSCAN, AgglomerativeClustering
from sklearn.metrics import silhouette_samples, silhouette_score
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.datasets import load_wine
from scipy.cluster.hierarchy import dendrogram, linkage
from matplotlib import cm
import seaborn as sns
from sklearn.feature_selection import SelectKBest, f_classif
from docx import Document
from docx.shared import Inches

# Get the current working directory
current_directory = os.getcwd()

# Define the project folder path dynamically within the current working directory
project_folder = os.path.join(current_directory, 'Small_Project_2')

# Create the project folder if it doesn't already exist
if not os.path.exists(project_folder):
    os.makedirs(project_folder)

# Load Wine dataset and remove labels
wine = load_wine()
wine_df = pd.DataFrame(wine.data, columns=wine.feature_names)

# Standardize the dataset
scaler = StandardScaler()
X_scaled = scaler.fit_transform(wine_df)

# Define a consistent colormap for clusters
cluster_cmap = cm.get_cmap("Set2")

# Function to visualize clustering results using PCA (for 2D projection)
def plot_clusters(X, labels, centroids=None, title="", filename=""):
    pca = PCA(n_components=2)
    X_pca = pca.fit_transform(X)
    unique_labels = np.unique(labels)
    
    # Use consistent colors for each cluster
    colors = cluster_cmap(np.linspace(0, 1, len(unique_labels)))
    
    for i, label in enumerate(unique_labels):
        plt.scatter(X_pca[labels == label, 0], X_pca[labels == label, 1], 
                    s=50, c=[colors[i]], label=f'Cluster {label+1}', edgecolor='k')
    
    if centroids is not None:
        centroids_pca = pca.transform(centroids)
        plt.scatter(centroids_pca[:, 0], centroids_pca[:, 1], s=250, marker='*', c='red', edgecolor='black', label='Centroids')
    
    plt.title(title)
    plt.legend()
    plt.grid(True)
    
    # Save the figure
    plt.savefig(os.path.join(project_folder, filename))
    plt.show()

### 1. K-means clustering with 3 clusters and k-means++ initialization ###
kmeans = KMeans(n_clusters=3, init='k-means++', n_init=10, random_state=0)
kmeans_labels = kmeans.fit_predict(X_scaled)

# Visualize K-means clusters with centroids and save the figure
plot_clusters(X_scaled, kmeans_labels, centroids=kmeans.cluster_centers_, title='K-means Clustering', filename='kmeans_clusters.png')

### 2. DBSCAN clustering ###
dbscan = DBSCAN(eps=0.5, min_samples=5)
dbscan_labels = dbscan.fit_predict(X_scaled)

# Visualize DBSCAN clusters and save the figure
plot_clusters(X_scaled, dbscan_labels, title='DBSCAN Clustering', filename='dbscan_clusters.png')

### 3. Hierarchical clustering ###
linked = linkage(X_scaled, method='complete')

# Visualize the hierarchical clustering as a dendrogram and save the figure
plt.figure(figsize=(10, 7))
dendrogram(linked)
plt.title('Hierarchical Clustering - Dendrogram')
plt.grid(True)
plt.savefig(os.path.join(project_folder, 'hierarchical_dendrogram.png'))
plt.show()

# Perform hierarchical clustering with 3 clusters and visualize
hierarchical = AgglomerativeClustering(n_clusters=3)
hierarchical_labels = hierarchical.fit_predict(X_scaled)
plot_clusters(X_scaled, hierarchical_labels, title='Hierarchical Clustering', filename='hierarchical_clusters.png')

### Feature selection: Select the top 5 features ###
selector = SelectKBest(f_classif, k=5)
X_selected = selector.fit_transform(X_scaled, wine.target)

# Get the selected feature names
selected_features = wine_df.columns[selector.get_support()]

# Perform hierarchical clustering on the selected features
linked_selected = linkage(X_selected.T, method='complete')

# Plot the dendrogram for the selected features and save the figure
plt.figure(figsize=(10, 7))
dendrogram(linked_selected, labels=selected_features)
plt.title('Dendrogram for Top Selected Features')
plt.grid(True)
plt.savefig(os.path.join(project_folder, 'dendrogram_selected_features.png'))
plt.show()

# Create a heatmap for the selected features and save the figure
plt.figure(figsize=(12, 8))
sns.heatmap(pd.DataFrame(X_selected, columns=selected_features).corr(), annot=True, cmap='coolwarm')
plt.title('Heatmap of Top Selected Features')
plt.savefig(os.path.join(project_folder, 'heatmap_selected_features.png'))
plt.show()

### 4. Silhouette analysis for K-means with 3, 4, and 5 clusters ###
def silhouette_analysis(X, range_n_clusters):
    for n_clusters in range_n_clusters:
        fig, (ax1, ax2) = plt.subplots(1, 2)
        fig.set_size_inches(18, 7)

        # Initialize the KMeans
        kmeans = KMeans(n_clusters=n_clusters, init='k-means++', n_init=10, random_state=42)
        cluster_labels = kmeans.fit_predict(X)

        # Silhouette score for the average value
        silhouette_avg = silhouette_score(X, cluster_labels)
        print(f"For n_clusters = {n_clusters}, the average silhouette score is {silhouette_avg}")

        # Compute the silhouette scores for each sample
        sample_silhouette_values = silhouette_samples(X, cluster_labels)

        y_lower = 10
        ax1.set_xlim([-0.1, 1])
        ax1.set_ylim([0, len(X) + (n_clusters + 1) * 10])

        # Use consistent colors for the silhouette plot
        colors = cluster_cmap(np.linspace(0, 1, n_clusters))

        for i in range(n_clusters):
            ith_cluster_silhouette_values = sample_silhouette_values[cluster_labels == i]
            ith_cluster_silhouette_values.sort()
            size_cluster_i = ith_cluster_silhouette_values.shape[0]
            y_upper = y_lower + size_cluster_i

            ax1.fill_betweenx(np.arange(y_lower, y_upper),
                              0, ith_cluster_silhouette_values,
                              facecolor=colors[i], edgecolor=colors[i], alpha=0.7)

            ax1.text(-0.05, y_lower + 0.5 * size_cluster_i, str(i))
            y_lower = y_upper + 10

        ax1.set_title(f"Silhouette plot for {n_clusters} clusters")
        ax1.set_xlabel("Silhouette coefficient values")
        ax1.set_ylabel("Cluster label")

        # Average silhouette line (axvline)
        ax1.axvline(x=silhouette_avg, color="red", linestyle="--")

        # Plot 2D clustering results using PCA
        pca = PCA(n_components=2)
        X_pca = pca.fit_transform(X)
        ax2.scatter(X_pca[:, 0], X_pca[:, 1], c=cluster_labels, cmap='nipy_spectral', marker='o', edgecolor='k', s=50)
        ax2.set_title(f"Cluster plot for {n_clusters} clusters")

        # Save the figure
        plt.savefig(os.path.join(project_folder, f'silhouette_{n_clusters}_clusters.png'))
        plt.show()

# Perform silhouette analysis for K-means with 3, 4, and 5 clusters
silhouette_analysis(X_scaled, [3, 4, 5])

### Generate Word Report ###
doc = Document()

# Add a title
doc.add_heading('Clustering Comparison Report', 0)

# Add K-means section
doc.add_heading('K-Means Clustering', level=1)
doc.add_paragraph('K-means clustering was performed with 3 clusters. The results, as visualized in the figure below, show distinct clusters with centroids marked.')
doc.add_picture(os.path.join(project_folder, 'kmeans_clusters.png'), width=Inches(5.0))

# Add DBSCAN section
doc.add_heading('DBSCAN Clustering', level=1)
doc.add_paragraph('DBSCAN clustering identified clusters based on density. The figure below demonstrates the identified clusters.')
doc.add_picture(os.path.join(project_folder, 'dbscan_clusters.png'), width=Inches(5.0))

# Add Hierarchical Clustering section
doc.add_heading('Hierarchical Clustering', level=1)
doc.add_paragraph('Hierarchical clustering was performed with 3 clusters. The dendrogram below illustrates the clustering hierarchy.')
doc.add_picture(os.path.join(project_folder, 'hierarchical_dendrogram.png'), width=Inches(5.0))

# Add silhouette analysis section
doc.add_heading('Silhouette Analysis', level=1)
doc.add_paragraph('Silhouette analysis for K-means clustering with different numbers of clusters (3, 4, and 5) was performed. The figures below show the silhouette plots.')
doc.add_picture(os.path.join(project_folder, 'silhouette_3_clusters.png'), width=Inches(5.0))
doc.add_picture(os.path.join(project_folder, 'silhouette_4_clusters.png'), width=Inches(5.0))
doc.add_picture(os.path.join(project_folder, 'silhouette_5_clusters.png'), width=Inches(5.0))

# Add a title
doc.add_heading('Clustering Algorithm Comparison Report', 0)

# Add Conclusion section
doc.add_heading('Conclusion', level=1)

# K-Means findings
doc.add_heading('K-Means Clustering:', level=2)
doc.add_paragraph('K-Means is effective for datasets with well-separated and compact clusters. In this case, it was able to partition '
                  'the Wine dataset into three clusters with good results. However, K-Means assumes the clusters are spherical and of equal size, '
                  'which might not always be true in real-world datasets.')

# DBSCAN findings
doc.add_heading('DBSCAN Clustering:', level=2)
doc.add_paragraph('DBSCAN was unable to produce meaningful clusters for the Wine dataset, likely due to the lack of clear density variations. '
                  'While DBSCAN is excellent for finding arbitrary-shaped clusters and handling noise, it requires careful parameter tuning '
                  'and may struggle in high-dimensional or densely packed datasets like Wine.')

# Hierarchical findings
doc.add_heading('Hierarchical Clustering:', level=2)
doc.add_paragraph('Hierarchical clustering effectively identified clusters in the Wine dataset and provided a clear visual representation of cluster relationships through the dendrogram. '
                  'It is useful when understanding the hierarchical structure between clusters is important, but it may become computationally expensive for larger datasets.')

# Recommendation for Real-World Use
doc.add_heading('Recommendation for Real-World Use:', level=2)
doc.add_paragraph('For datasets like the Wine dataset, K-Means is recommended due to its simplicity and effectiveness at finding clear, well-separated clusters. '
                  'Hierarchical clustering is useful when insight into the relationships between clusters is needed. DBSCAN should be considered for datasets with noise or when irregular cluster shapes are expected, but it may struggle without the right density parameters.')

# Save the document in the project folder
report_path = os.path.join(project_folder, 'Clustering_Comparison_Report.docx')
doc.save(report_path)

print("Clustering report generated and saved successfully.")

print(f"Report generated and saved at {report_path}")
